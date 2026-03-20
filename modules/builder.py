import os
import re
import docx
from docx.shared import Pt, Inches
from .extractor import load_state


def build_docx_from_markdown(state_file, output_path):
    """
    Builds a Word document from translated Markdown sections stored in state.
    Parses basic Markdown (headings, bold, tables, paragraphs) into python-docx elements.
    """
    state = load_state(state_file)
    
    # Collect translated sections in order
    section_keys = sorted(
        [k for k in state if k.startswith("section_")],
        key=lambda k: int(k.split("_")[1])
    )
    
    translated_count = sum(
        1 for k in section_keys if state[k]["status"] == "translated"
    )
    print(f"[Builder] Building DOCX from {translated_count}/{len(section_keys)} translated sections.")
    
    doc = docx.Document()
    
    # Set default font for the document
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Noto Sans JP"
    font.size = Pt(10.5)
    
    for key in section_keys:
        block = state[key]
        if block["status"] == "translated":
            text = block["translated_text"]
        else:
            # Fall back to original text if not yet translated
            text = block["original_text"]
        
        render_markdown_section(doc, text)
    
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    print(f"[Builder] Successfully built translated Word document at {output_path}")
    return True


def render_markdown_section(doc, md_text):
    """
    Renders a Markdown section into the docx document.
    Handles: headings, bold text, tables, and plain paragraphs.
    """
    lines = md_text.split("\n")
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Image: ![alt](path)
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if img_match:
            img_path = img_match.group(2)
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(4.5))
                except Exception as e:
                    print(f"Warning: Failed to embed image {img_path}: {e}")
                    doc.add_paragraph(f"[Image: {img_path}]")
            else:
                doc.add_paragraph(f"[Image not found: {img_path}]")
            i += 1
            continue
        
        # Heading
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            # Remove ** from heading text if present
            heading_text = re.sub(r'\*\*(.*?)\*\*', r'\1', heading_text)
            p = doc.add_heading(heading_text, level=min(level, 4))
            i += 1
            continue
        
        # Table: detect table rows starting with |
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            render_table(doc, table_lines)
            continue
        
        # Regular paragraph: collect consecutive non-special lines
        para_lines = []
        while i < len(lines):
            current = lines[i]
            if not current.strip():
                i += 1
                break
            if re.match(r'^#{1,6}\s', current):
                break
            if current.strip().startswith("|"):
                break
            if re.match(r'^!\[', current.strip()):
                break
            para_lines.append(current.strip())
            i += 1
        
        if para_lines:
            para_text = " ".join(para_lines)
            p = doc.add_paragraph()
            render_inline_markdown(p, para_text)


def render_inline_markdown(paragraph, text):
    """
    Renders inline Markdown (bold, italic, links) into a paragraph using runs.
    """
    # Split text by **bold** markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Regular text — also handle *italic*
            italic_parts = re.split(r'(\*[^*]+?\*)', part)
            for ip in italic_parts:
                if ip.startswith("*") and ip.endswith("*") and not ip.startswith("**"):
                    run = paragraph.add_run(ip[1:-1])
                    run.italic = True
                else:
                    if ip:
                        paragraph.add_run(ip)


def render_table(doc, table_lines):
    """
    Renders a Markdown table into a docx table.
    Expects lines like: | col1 | col2 | col3 |
    Second line is separator: | --- | --- | --- |
    """
    if len(table_lines) < 2:
        # Not a valid table, render as paragraphs
        for line in table_lines:
            doc.add_paragraph(line)
        return
    
    # Parse header
    header_cells = parse_table_row(table_lines[0])
    
    # Check if second line is separator
    separator_idx = 1
    if re.match(r'\|[\s\-:]+\|', table_lines[1]):
        separator_idx = 2  # Skip separator line
    
    # Parse data rows
    data_rows = []
    for line in table_lines[separator_idx:]:
        cells = parse_table_row(line)
        if cells:
            data_rows.append(cells)
    
    if not header_cells:
        return
    
    num_cols = len(header_cells)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = "Table Grid"
    
    # Fill header
    for j, cell_text in enumerate(header_cells):
        if j < num_cols:
            cell = table.rows[0].cells[j]
            cell.text = cell_text
            # Bold header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    
    # Fill data rows
    for row_idx, row_data in enumerate(data_rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                table.rows[row_idx + 1].cells[j].text = cell_text


def parse_table_row(line):
    """Parse a markdown table row into a list of cell texts."""
    # Remove leading/trailing |
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    
    cells = [cell.strip() for cell in line.split("|")]
    return cells
