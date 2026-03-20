import fitz  # PyMuPDF
import json
import os
import re
import docx


def load_state(state_file):
    if os.path.exists(state_file):
        with open(state_file, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_state(state, state_file):
    with open(state_file, "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=2)

def extract_pdf_blocks(pdf_path, state_file, limit=None):
    """
    Extracts text blocks from the PDF and updates the state file.
    Only considers blocks of type 0 (text).
    """
    doc = fitz.open(pdf_path)
    state = load_state(state_file)
    
    num_pages = min(len(doc), limit) if limit else len(doc)
    
    # Track which IDs are present in this run
    current_ids = set()

    for page_num in range(num_pages):
        page = doc[page_num]
        blocks = page.get_text("dict")["blocks"]
        
        for block_idx, block in enumerate(blocks):
            if block.get("type") == 0:  # text block
                # Reconstruct full block text
                block_text = ""
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        block_text += span.get("text", "")
                    block_text += "\n"
                
                # Safely merge hyphenated words across lines
                block_text = re.sub(r'([a-zA-Z])-\n([a-zA-Z])', r'\1\2', block_text)
                block_text = block_text.strip()
                if len(block_text) < 2:
                    continue
                
                block_id = f"p{page_num}_b{block_idx}"
                current_ids.add(block_id)
                
                if block_id not in state:
                    state[block_id] = {
                        "id": block_id,
                        "page": page_num,
                        "bbox": block["bbox"],  # [x0, y0, x1, y1]
                        "original_text": block_text,
                        "translated_text": "",
                        "status": "pending",
                        "retry_count": 0
                    }
                else:
                    # Update coordinates just in case
                    state[block_id]["bbox"] = block["bbox"]
                    state[block_id]["original_text"] = block_text

    save_state(state, state_file)
    print(f"[Extractor] Extracted/Updated {len(current_ids)} text blocks from {num_pages} pages.")
    return state

def extract_docx_blocks(docx_path, state_file):
    """
    Extracts text paragraphs from the docx file (both body and tables),
    and updates the state file.
    """
    doc = docx.Document(docx_path)
    state = load_state(state_file)
    
    current_ids = set()
    
    def process_paragraphs(paragraphs, prefix):
        for i, p in enumerate(paragraphs):
            text = p.text.strip()
            if len(text) < 2:
                continue
            
            block_id = f"{prefix}_p{i}"
            current_ids.add(block_id)
            
            if block_id not in state:
                state[block_id] = {
                    "id": block_id,
                    "original_text": text,
                    "translated_text": "",
                    "status": "pending",
                    "retry_count": 0
                }
            else:
                state[block_id]["original_text"] = text

    # Process body paragraphs
    process_paragraphs(doc.paragraphs, "body")
    
    # Process table paragraphs
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                prefix = f"t{t_idx}_r{r_idx}_c{c_idx}"
                process_paragraphs(cell.paragraphs, prefix)
                
    save_state(state, state_file)
    print(f"[Extractor] Extracted/Updated {len(current_ids)} text blocks from Word document.")
    return state
